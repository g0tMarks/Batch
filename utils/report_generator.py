from langchain.prompts import PromptTemplate
from langchain_anthropic import ChatAnthropic
from langchain.schema import AIMessage
from docx import Document
import time
import os
import asyncio
from typing import List, Dict, Any
from datetime import datetime
from dotenv import load_dotenv
import logging

# Set up logging
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

class ReportGenerationError(Exception):
    """Base exception for report generation errors"""
    pass

class ReportGenerationService:
    def __init__(self):
        api_key = os.getenv('ANTHROPIC_API_KEY')
        if not api_key:
            raise ReportGenerationError("ANTHROPIC_API_KEY environment variable is not set")
            
        logger.debug("Initializing ReportGenerationService")
        self.llm = ChatAnthropic(
            model="claude-3-opus-20240229",
            anthropic_api_key=api_key,
            temperature=0.4
        )
        self.prompt_template = PromptTemplate(
            input_variables=["student_name", "year", "gender", "academic_performance", 
                           "extracurricular_activities", "other", "sample_report"],
            template="""You are an experienced and caring high school teacher, your job is to write a report for students in your Home Group that comments on their academic performance, their wellbeing and their involvement in extracurricular activities.

Use this sample report below as a template.

{sample_report}

It is very important that you follow the template above, using the same structure, tone, language and length. You can vary adjectives in the closing sentence but keep the opening sentence the same. It must be written in Australian English.

Write a report for {student_name}.

They are in year {year}, use this year and class in place of any other mentions of year and class, such as 7A, 9B, 10D etc.

Their gender is {gender}, use the appropriate pro-nouns.

Academic Performance:
{academic_performance}

Extracurricular Activities:
For House Athletics and House swimming, do not list out all the events the student participate in, just provide an overview of their involvement with a general comment on the events they participated in.
{extracurricular_activities}

Other important information to include 1 sentence on:
{other}"""
        )
        logger.debug("ReportGenerationService initialized successfully")

    def generate_prompt(self, student: Dict[str, Any]) -> str:
        """Generate a prompt for a single student"""
        try:
            logger.debug(f"Generating prompt for student: {student.get('student_name', 'unknown')}")
            prompt = self.prompt_template.format(**student)
            logger.debug(f"Generated prompt: {prompt[:200]}...")  # Log first 200 chars of prompt
            return prompt
        except KeyError as e:
            logger.error(f"Missing required student data: {str(e)}")
            raise ReportGenerationError(f"Missing required student data: {str(e)}")
        except Exception as e:
            logger.error(f"Error generating prompt: {str(e)}")
            raise ReportGenerationError(f"Error generating prompt: {str(e)}")

    async def generate_single_report(self, student: Dict[str, Any]) -> str:
        """Generate a report for a single student"""
        try:
            logger.debug(f"Starting report generation for student: {student.get('student_name', 'unknown')}")
            prompt = self.generate_prompt(student)
            
            # Use the LLM to generate the report
            response = await self.llm.ainvoke(prompt)
            if isinstance(response, AIMessage):
                report = response.content
            else:
                report = str(response)
                
            logger.debug(f"Generated report for {student.get('student_name', 'unknown')}")
            return report
        except Exception as e:
            logger.error(f"Error generating report for {student.get('student_name', 'unknown')}: {str(e)}")
            raise ReportGenerationError(f"Error generating report for {student.get('student_name', 'unknown')}: {str(e)}")

    async def generate_reports(self, student_list: List[Dict[str, Any]]) -> List[str]:
        """Generate reports for multiple students concurrently"""
        try:
            logger.info(f"Starting batch report generation for {len(student_list)} students")
            reports = []
            for student in student_list:
                report = await self.generate_single_report(student)
                reports.append(report)
            logger.info(f"Successfully generated {len(reports)} reports")
            return reports
        except Exception as e:
            logger.error(f"Error generating batch reports: {str(e)}")
            raise ReportGenerationError(f"Error generating batch reports: {str(e)}")

    async def generate_reports_with_progress(self, student_list: List[Dict[str, Any]], user_id: str, progress_tracker: dict) -> List[str]:
        """Generate reports for multiple students with progress tracking"""
        try:
            logger.info(f"Starting batch report generation for {len(student_list)} students")
            reports = []
            for i, student in enumerate(student_list, 1):
                # Update progress
                progress_tracker[user_id] = {
                    'current': i,
                    'total': len(student_list),
                    'status': f'Generating report {i}/{len(student_list)}',
                    'progress': int((i / len(student_list)) * 90)  # Reserve 10% for final steps
                }
                
                report = await self.generate_single_report(student)
                reports.append(report)
                logger.debug(f"Generated report {i}/{len(student_list)}")
                
            logger.info(f"Successfully generated {len(reports)} reports")
            return reports
        except Exception as e:
            logger.error(f"Error generating batch reports: {str(e)}")
            raise ReportGenerationError(f"Error generating batch reports: {str(e)}")

    async def create_word_doc(self, reports: List[str], output_dir: str = "/tmp/reports") -> str:
        """Create a Word document containing all reports"""
        try:
            logger.debug(f"Creating Word document with {len(reports)} reports")
            os.makedirs(output_dir, exist_ok=True)
            doc = Document()
            
            for i, report in enumerate(reports, 1):
                logger.debug(f"Adding report {i} to document")
                doc.add_heading(f"Student Report {i}", level=1)
                doc.add_paragraph(report)
                if i < len(reports):
                    doc.add_page_break()
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(output_dir, f"student_reports_{timestamp}.docx")
            logger.debug(f"Saving document to: {output_path}")
            doc.save(output_path)
            logger.info(f"Successfully created Word document at {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error creating Word document: {str(e)}")
            raise ReportGenerationError(f"Error creating Word document: {str(e)}")
